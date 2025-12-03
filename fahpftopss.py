# ===========================================================
#  PFS–FAHP + PFS–FTOPSIS TKINTER GUI
#  - Modern Soft Tema
#  - Scrollable Sekmeler (FAHP/FTOPSIS)
#  - Renkli FAHP Ağırlık Tablosu
#  - FTOPSIS Grafik
#  - PDF & Word Rapor (Konum Seçmeli)
#  - "Hazırlayan: Ahmet Esad Boydak" Butonu
# ===========================================================

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pandas as pd
import math
from dataclasses import dataclass
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from datetime import datetime
import json
import os


# ===========================================================
# =============== 1. BÖLÜM: SCROLLABLE FRAME ================
# ===========================================================

class ScrollableFrame(ttk.Frame):
    """
    Notebook sekmeleri içinde scroll (kaydırma) desteği için
    Canvas + Frame + Scrollbar kombinasyonu
    """
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)

        canvas = tk.Canvas(self, borderwidth=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")


# ===========================================================
# === 2. BÖLÜM: PFS SINIFI ve DILSEL ÖLÇEK TANIMLARI =========
# ===========================================================

@dataclass
class PFS:
    mu: float
    nu: float

    def __post_init__(self):
        if self.mu**2 + self.nu**2 > 1 + 1e-9:
            raise ValueError("Pythagorean condition violated: mu^2 + nu^2 <= 1")

    def score(self):
        return self.mu**2 - self.nu**2


# FAHP dilsel ölçek
AHP_SCALE = {
    "EQ": (0.70, 0.70),
    "LI": (0.30, 0.80),
    "MI": (0.50, 0.70),
    "HI": (0.70, 0.50),
    "VHI": (0.90, 0.30)
}

# FTOPSIS dilsel ölçek
TOPSIS_SCALE = {
    "VL": (0.10, 0.90),
    "L":  (0.30, 0.80),
    "M":  (0.50, 0.70),
    "H":  (0.70, 0.50),
    "VH": (0.90, 0.30)
}


# ===========================================================
# ======= 3. BÖLÜM: FAHP ve FTOPSIS HESAP FONKSIYONLARI ======
# ===========================================================

def fahp_weights(pfs_matrix):
    """
    pfs_matrix: (5x5) birleşik PFS matrisi (aggregate FAHP)
    Dönen: 5 elemanlı ağırlık listesi
    """
    n = len(pfs_matrix)
    S = [[pfs_matrix[i][j].mu for j in range(n)] for i in range(n)]

    g = []
    for i in range(n):
        prod = 1.0
        for j in range(n):
            prod *= S[i][j]
        g.append(prod ** (1.0 / n))

    total = sum(g)
    if total == 0:
        return [1.0 / n] * n
    return [gi / total for gi in g]


def ftopsis(pfs_matrix, weights, alternatives):
    """
    pfs_matrix: m x n PFS matrisi (aggregate TOPSIS)
    weights: n boyutlu ağırlık vektörü
    alternatives: ['B1', 'B2', ...]
    Dönen: sıralanmış DataFrame
    """
    m = len(pfs_matrix)
    n = len(pfs_matrix[0])

    # Skor matrisi
    S = [[pfs_matrix[i][j].score() for j in range(n)] for i in range(m)]

    # Normalize
    R = []
    for j in range(n):
        denom = math.sqrt(sum(S[i][j] ** 2 for i in range(m)))
        if denom == 0:
            denom = 1.0
        col = [S[i][j] / denom for i in range(m)]
        if j == 0:
            R = [[col[i]] for i in range(m)]
        else:
            for i in range(m):
                R[i].append(col[i])

    # Ağırlıklı normalize matris
    V = [[R[i][j] * weights[j] for j in range(n)] for i in range(m)]

    # Pozitif / negatif ideal
    v_plus = [max(V[i][j] for i in range(m)) for j in range(n)]
    v_minus = [min(V[i][j] for i in range(m)) for j in range(n)]

    # Uzaklıklar
    D_plus, D_minus = [], []
    for i in range(m):
        dp = math.sqrt(sum((V[i][j] - v_plus[j]) ** 2 for j in range(n)))
        dm = math.sqrt(sum((V[i][j] - v_minus[j]) ** 2 for j in range(n)))
        D_plus.append(dp)
        D_minus.append(dm)

    CC = [
        D_minus[i] / (D_plus[i] + D_minus[i]) if (D_plus[i] + D_minus[i]) != 0 else 0.0
        for i in range(m)
    ]

    df = pd.DataFrame({
        "Alternative": alternatives,
        "D+": D_plus,
        "D-": D_minus,
        "CC": CC
    }).sort_values(by="CC", ascending=False).reset_index(drop=True)

    return df


# ===========================================================
# =============== 8. BÖLÜM: SENARYO YÖNETİMİ ==================
# ===========================================================

class ScenarioManager:
    """Senaryo kaydetme ve karşılaştırması"""
    def __init__(self, data_dir="scenarios"):
        self.data_dir = data_dir
        if not os.path.exists(data_dir):
            os.makedirs(data_dir)
    
    def save_scenario(self, scenario_name, fahp_data, ftopsis_data, weights, result_df):
        data = {
            "name": scenario_name,
            "date": datetime.now().isoformat(),
            "fahp_ling": fahp_data,
            "ftopsis_ling": ftopsis_data,
            "weights": weights,
            "results": result_df.to_dict()
        }
        path = os.path.join(self.data_dir, f"{scenario_name}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return path
    
    def load_scenario(self, scenario_name):
        path = os.path.join(self.data_dir, f"{scenario_name}.json")
        if not os.path.exists(path):
            return None
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    
    def list_scenarios(self):
        if not os.path.exists(self.data_dir):
            return []
        files = [f.replace(".json", "") for f in os.listdir(self.data_dir) if f.endswith(".json")]
        return sorted(files)


# ===========================================================
# =============== 9. BÖLÜM: İSTATİSTİK ANALİZİ ===============
# ===========================================================

def statistical_summary(weights, result_df):
    """İstatistiksel özet bilgisi"""
    cc_values = result_df["CC"].values
    
    summary = {
        "Ortalama CC": float(cc_values.mean()),
        "Standart Sapma": float(cc_values.std()),
        "Min CC": float(cc_values.min()),
        "Max CC": float(cc_values.max()),
        "CC Aralığı": float(cc_values.max() - cc_values.min()),
        "Ağırlık Ortalaması": float(sum(weights) / len(weights)),
        "Max Ağırlık": float(max(weights)),
        "Min Ağırlık": float(min(weights))
    }
    return summary


# ===========================================================
# =============== 10. BÖLÜM: UZMAN UYUŞUM ANALİZİ ============
# ===========================================================

def expert_agreement_analysis(pfs_experts):
    """Uzmanlar arası uyum analizi"""
    if not pfs_experts or not pfs_experts[0]:
        return {"Ortalama Varyans": 0, "Uyum Skoru": 1, "Uyum Seviyesi": "Yüksek"}
    
    n = len(pfs_experts[0])
    m = len(pfs_experts[0][0])
    variances = []
    
    for i in range(n):
        for j in range(m):
            mu_values = [pfs_experts[e][i][j].mu for e in range(len(pfs_experts))]
            nu_values = [pfs_experts[e][i][j].nu for e in range(len(pfs_experts))]
            mu_avg = sum(mu_values) / len(mu_values)
            nu_avg = sum(nu_values) / len(nu_values)
            var = (sum((mu - mu_avg)**2 for mu in mu_values) / len(mu_values)) + \
                  (sum((nu - nu_avg)**2 for nu in nu_values) / len(nu_values))
            variances.append(var)
    
    avg_variance = sum(variances) / len(variances) if variances else 0
    return {
        "Ortalama Varyans": round(avg_variance, 4),
        "Uyum Skoru": round(1 - min(avg_variance, 1), 4),
        "Uyum Seviyesi": "Yüksek" if avg_variance < 0.1 else "Orta" if avg_variance < 0.2 else "Düşük"
    }


# ===========================================================
# =============== 11. BÖLÜM: RADAR CHART =====================
# ===========================================================

def create_radar_chart(criteria, weights):
    """Kriterler için radar chart"""
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(projection='polar'))
    
    angles = [n / float(len(criteria)) * 2 * math.pi for n in range(len(criteria))]
    weights_plot = list(weights) + [weights[0]]
    angles += angles[:1]
    
    ax.plot(angles, weights_plot, 'o-', linewidth=2, color="#5680e9")
    ax.fill(angles, weights_plot, alpha=0.25, color="#5680e9")
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(criteria)
    ax.set_ylim(0, max(weights) * 1.1)
    ax.set_title("Kriter Ağırlıkları (Radar Chart)", pad=20)
    ax.grid(True)
    
    return fig


# ===========================================================
# =============== 4. BÖLÜM: ANA GUI SINIFI ===================
# ===========================================================

class App:

    def __init__(self, root):
        self.root = root
        self.root.title("PFS–FAHP + PFS–FTOPSIS (AEB)")
        self.root.geometry("1400x800")
        self.root.configure(bg="#f2f2f2")

        self.criteria = ["A1", "A2", "A3", "A4", "A5"]
        self.alternatives = ["B1", "B2", "B3", "B4", "B5"]
        self.experts = 4
        
        self.scenario_manager = ScenarioManager()
        self.last_result = None
        self.current_theme = "light"

        # Header Frame
        header_frame = ttk.Frame(root)
        header_frame.pack(pady=15)



        # Başlık
        header = ttk.Label(
            header_frame,
            text="Bulanık Çok Kriterli Karar Verme:\nFAHP & FTOPSIS Analizi",
            font=("Segoe UI", 22, "bold"),
            background="#f2f2f2",
            foreground="#2c3e50"
        )
        header.pack(side="left", padx=10)

        
        

        self.tabs = ttk.Notebook(root)
        self.tabs.pack(expand=1, fill="both", padx=20, pady=10)

        self.fahp_frames = []
        self.ftopsis_frames = []

        self.create_fahp_tabs()
        self.create_ftopsis_tabs()

        ttk.Button(root, text="HESAPLA", command=self.compute, width=20).pack(pady=10)

        # Buton Frame
        btn_frame = ttk.Frame(root)
        btn_frame.pack(pady=10)

        ttk.Button(btn_frame, text="Tema Değiştir", command=self.toggle_theme).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="Senaryo Kaydet", command=self.save_scenario_dialog).grid(row=0, column=1, padx=5)
        ttk.Button(btn_frame, text="Senaryo Yükle", command=self.load_scenario_dialog).grid(row=0, column=2, padx=5)
        ttk.Button(btn_frame, text="Senaryoları Karşılaştır", command=self.compare_scenarios).grid(row=0, column=3, padx=5)

        info_btn = ttk.Button(root, text="Hazırlayan", command=self.show_info)
        info_btn.place(relx=0.97, rely=0.97, anchor="se")
        scale_btn = ttk.Button(root, text="Dilsel Ölçek", command=self.show_scales)
        scale_btn.place(relx=0.84, rely=0.97, anchor="se")

    def toggle_theme(self):
        """Açık/Koyu tema geçişi"""
        if self.current_theme == "light":
            self.root.configure(bg="#2b2b2b")
            self.current_theme = "dark"
            messagebox.showinfo("Tema", "Koyu tema etkinleştirildi")
        else:
            self.root.configure(bg="#f2f2f2")
            self.current_theme = "light"
            messagebox.showinfo("Tema", "Açık tema etkinleştirildi")

    def save_scenario_dialog(self):
        """Senaryo kaydetme dialogu"""
        if self.last_result is None:
            messagebox.showwarning("Uyarı", "Önce HESAPLA butonuna tıklayın")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Senaryo Kaydet")
        dialog.geometry("400x150")
        
        ttk.Label(dialog, text="Senaryo Adı:").pack(pady=10)
        entry = ttk.Entry(dialog, width=30)
        entry.pack(pady=5)
        
        def save():
            name = entry.get().strip()
            if not name:
                messagebox.showwarning("Hata", "Senaryo adı boş olamaz")
                return
            
            self.scenario_manager.save_scenario(
                name,
                self.last_result["fahp_ling"],
                self.last_result["ftopsis_ling"],
                self.last_result["weights"],
                self.last_result["df"]
            )
            messagebox.showinfo("Başarılı", f"Senaryo '{name}' kaydedildi")
            dialog.destroy()
        
        ttk.Button(dialog, text="Kaydet", command=save).pack(pady=10)

    def load_scenario_dialog(self):
        """Senaryo yükleme dialogu"""
        scenarios = self.scenario_manager.list_scenarios()
        if not scenarios:
            messagebox.showinfo("Bilgi", "Kaydedilmiş senaryo yok")
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Senaryo Yükle")
        dialog.geometry("400x300")
        
        ttk.Label(dialog, text="Senaryolar:").pack(pady=10)
        
        listbox = tk.Listbox(dialog, height=10)
        listbox.pack(fill="both", expand=True, padx=10, pady=5)
        
        for scenario in scenarios:
            listbox.insert(tk.END, scenario)
        
        def load():
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("Hata", "Senaryo seçin")
                return
            
            scenario_name = scenarios[selection[0]]
            data = self.scenario_manager.load_scenario(scenario_name)
            
            messagebox.showinfo("Başarılı", 
                f"Senaryo '{scenario_name}' yüklendi\n"
                f"Tarih: {data['date']}\n"
                f"Ağırlıklar: {', '.join(f'{w:.4f}' for w in data['weights'])}"
            )
            
            # Veriyi girişlere yazdır
            self.restore_scenario_to_inputs(data)
            dialog.destroy()
        
        def load_and_restore():
            selection = listbox.curselection()
            if not selection:
                messagebox.showwarning("Hata", "Senaryo seçin")
                return
            
            scenario_name = scenarios[selection[0]]
            data = self.scenario_manager.load_scenario(scenario_name)
            
            # Veriyi girişlere yazdır
            self.restore_scenario_to_inputs(data)
            
            messagebox.showinfo("Başarılı", 
                f"Senaryo '{scenario_name}' girişlere yüklendi!"
            )
            dialog.destroy()
        
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Sadece Bilgi Göster", command=load).grid(row=0, column=0, padx=5)
        ttk.Button(btn_frame, text="Girişlere Yükle", command=load_and_restore).grid(row=0, column=1, padx=5)

    def compare_scenarios(self):
        """Senaryo karşılaştırması"""
        scenarios = self.scenario_manager.list_scenarios()
        if len(scenarios) < 2:
            messagebox.showinfo("Bilgi", "Karşılaştırma için en az 2 senaryo gerekli")
            return
        
        win = tk.Toplevel(self.root)
        win.title("Senaryo Karşılaştırması")
        win.geometry("600x400")
        
        txt = tk.Text(win, font=("Segoe UI", 11))
        txt.pack(fill="both", expand=True, padx=10, pady=10)
        
        for scenario_name in scenarios:
            data = self.scenario_manager.load_scenario(scenario_name)
            txt.insert("end", f"\n{'='*50}\n")
            txt.insert("end", f"Senaryo: {scenario_name}\n")
            txt.insert("end", f"Tarih: {data['date']}\n")
            txt.insert("end", f"Ağırlıklar: {', '.join(f'{w:.4f}' for w in data['weights'])}\n")
            txt.insert("end", f"{'='*50}\n\n")
        
        txt.configure(state="disabled")

    def restore_scenario_to_inputs(self, data):
        """Kaydedilmiş senaryoyu girişlere geri yazdırır"""
        fahp_ling = data["fahp_ling"]
        ftopsis_ling = data["ftopsis_ling"]
        
        try:
            # FAHP sekmeleri güncelle
            for e in range(self.experts):
                for i in range(5):
                    for j in range(5):
                        w = self.fahp_frames[e][i][j]
                        value = fahp_ling[e][i][j]
                        
                        # Diagonal: string ("EQ"), değiştirilmez
                        if isinstance(w, str):
                            continue
                        # Combobox widget ise
                        else:
                            try:
                                w.set(value)
                            except:
                                pass
            
            # FTOPSIS sekmeleri güncelle
            for e in range(self.experts):
                for i in range(5):
                    for j in range(5):
                        w = self.ftopsis_frames[e][i][j]
                        value = ftopsis_ling[e][i][j]
                        
                        # Combobox widget ise
                        try:
                            w.set(value)
                        except:
                            pass
            
            messagebox.showinfo("Başarılı", "Tüm girişler senaryo verileriyle güncellendi!")
        
        except Exception as err:
            messagebox.showerror("Hata", f"Girişler yüklenirken hata:\n{str(err)}\n\nLütfen verilerin uyumlu olduğundan emin olun.")

    # ------------------ Hakkında ---------------------
    def show_info(self):
        messagebox.showinfo("Hazırlayan", "İsim: Ahmet Esad\nSoyad: Boydak")

    def show_scales(self):
        win = tk.Toplevel(self.root)
        win.title("Dilsel İfadeler ve PFS Karşılıkları")
        win.geometry("700x500")
        win.configure(bg="#f0f0f0")

        ttk.Label(
            win,
            text="FAHP – Dilsel Ölçek Tablosu",
            font=("Segoe UI", 14, "bold")
        ).pack(pady=10)

        # FAHP tablosu
        fahp_table = tk.Text(win, height=10, font=("Segoe UI", 11))
        fahp_table.pack(fill="x", padx=20)
        fahp_table.insert("end",
                         "İfade\t(μ, ν)\n"
                         "----------------------------------\n"
                         "EQ eşit önemli\t(0.70, 0.70)\n"
                         "LI az önemli\t(0.30, 0.80)\n"
                         "MI orta önemli\t(0.50, 0.70)\n"
                         "HI fazla önemli\t(0.70, 0.50)\n"
                         "VHI çok fazla önemli\t(0.90, 0.30)\n"
                         )
        fahp_table.configure(state="disabled")

        ttk.Label(
            win,
            text="FTOPSIS – Dilsel Ölçek Tablosu",
            font=("Segoe UI", 14, "bold")
        ).pack(pady=10)

        # FTOPSIS tablosu
        ft_table = tk.Text(win, height=10, font=("Segoe UI", 11))
        ft_table.pack(fill="x", padx=20)
        ft_table.insert("end",
                        "İfade\t(μ, ν)\n"
                        "----------------------------------\n"
                        "VL çok düşük uygun\t(0.10, 0.90)\n"
                        "L düşük uygun\t(0.30, 0.80)\n"
                        "M orta uygun\t(0.50, 0.70)\n"
                        "H yüksek uygun\t(0.70, 0.50)\n"
                        "VH çok yüksek uygun\t(0.90, 0.30)\n"
                        )
        ft_table.configure(state="disabled")

    # ------------------ FAHP Sekmeleri ---------------
    def create_fahp_tabs(self):
        for e in range(self.experts):
            outer = ScrollableFrame(self.tabs)
            frame = outer.scrollable_frame
            self.tabs.add(outer, text=f"FAHP Uzman {e+1}")

            combos = []

            ttk.Label(frame, text="", width=10).grid(row=0, column=0)
            for j, crit in enumerate(self.criteria):
                ttk.Label(frame, text=crit, font=("Segoe UI", 10, "bold")).grid(row=0, column=j+1, pady=8)

            for i in range(5):
                row = []
                ttk.Label(frame, text=self.criteria[i]).grid(row=i+1, column=0)

                for j in range(5):
                    if i == j:
                        lbl = ttk.Label(frame, text="EQ")
                        lbl.grid(row=i+1, column=j+1)
                        row.append("EQ")  # direkt string
                    else:
                        cb = ttk.Combobox(
                            frame,
                            values=list(AHP_SCALE.keys()),
                            width=7,
                            state="readonly"
                        )
                        cb.set("EQ")
                        cb.grid(row=i+1, column=j+1, padx=5, pady=4)
                        row.append(cb)
                combos.append(row)

            self.fahp_frames.append(combos)

    # ------------------ FTOPSIS Sekmeleri ------------
    def create_ftopsis_tabs(self):
        for e in range(self.experts):
            outer = ScrollableFrame(self.tabs)
            frame = outer.scrollable_frame
            self.tabs.add(outer, text=f"FTOPSIS Uzman {e+1}")

            combos = []

            ttk.Label(frame, text="", width=10).grid(row=0, column=0)
            for j, crit in enumerate(self.criteria):
                ttk.Label(frame, text=crit, font=("Segoe UI", 10, "bold")).grid(row=0, column=j+1, pady=8)

            for i in range(5):
                row = []
                ttk.Label(frame, text=self.alternatives[i]).grid(row=i+1, column=0)

                for j in range(5):
                    cb = ttk.Combobox(
                        frame,
                        values=list(TOPSIS_SCALE.keys()),
                        width=7,
                        state="readonly"
                    )
                    cb.set("M")
                    cb.grid(row=i+1, column=j+1, padx=5, pady=4)
                    row.append(cb)
                combos.append(row)

            self.ftopsis_frames.append(combos)

    # =================================================
    # 5. BÖLÜM: HESAPLAMA & SONUÇ PENCERESİ
    # =================================================
    def compute(self):
        # -----------------------------
        # FAHP – 4 uzman dilsel matris
        # -----------------------------
        fahp_ling_experts = []
        for e in range(self.experts):
            ling_matrix = []
            for i in range(5):
                row_labels = []
                for j in range(5):
                    if i == j:
                        label = "EQ"
                    else:
                        w = self.fahp_frames[e][i][j]
                        label = w if isinstance(w, str) else w.get()
                    row_labels.append(label)
                ling_matrix.append(row_labels)
            fahp_ling_experts.append(ling_matrix)

        # Dilsel → PFS (uzman bazında)
        fahp_pfs_experts = []
        for exp in fahp_ling_experts:
            matrix = []
            for i in range(5):
                row = []
                for j in range(5):
                    mu, nu = AHP_SCALE[exp[i][j]]
                    row.append(PFS(mu, nu))
                matrix.append(row)
            fahp_pfs_experts.append(matrix)

        # Uzmanları birleştir (ortalama mu, nu)
        pfs_fahp = []
        for i in range(5):
            row = []
            for j in range(5):
                mu_avg = sum(fahp_pfs_experts[e][i][j].mu for e in range(self.experts)) / self.experts
                nu_avg = sum(fahp_pfs_experts[e][i][j].nu for e in range(self.experts)) / self.experts
                row.append(PFS(mu_avg, nu_avg))
            pfs_fahp.append(row)

        # FAHP ağırlıkları
        weights = fahp_weights(pfs_fahp)

        # -----------------------------
        # FTOPSIS – 4 uzman dilsel matris
        # -----------------------------
        ftopsis_ling_experts = []
        for e in range(self.experts):
            ling_matrix = []
            for i in range(5):
                row_labels = []
                for j in range(5):
                    w = self.ftopsis_frames[e][i][j]
                    label = w.get()
                    row_labels.append(label)
                ling_matrix.append(row_labels)
            ftopsis_ling_experts.append(ling_matrix)

        # Dilsel → PFS (uzman bazında)
        ftopsis_pfs_experts = []
        for exp in ftopsis_ling_experts:
            matrix = []
            for i in range(5):
                row = []
                for j in range(5):
                    mu, nu = TOPSIS_SCALE[exp[i][j]]
                    row.append(PFS(mu, nu))
                matrix.append(row)
            ftopsis_pfs_experts.append(matrix)

        # Uzmanları birleştir (ortalama mu, nu)
        pfs_top = []
        for i in range(5):
            row = []
            for j in range(5):
                mu_avg = sum(ftopsis_pfs_experts[e][i][j].mu for e in range(self.experts)) / self.experts
                nu_avg = sum(ftopsis_pfs_experts[e][i][j].nu for e in range(self.experts)) / self.experts
                row.append(PFS(mu_avg, nu_avg))
            pfs_top.append(row)

        # FTOPSIS sonuçları
        result_df = ftopsis(pfs_top, weights, self.alternatives)

        # Grafik kaydet
        fig, ax = plt.subplots(figsize=(6, 4), dpi=100)
        ax.bar(result_df["Alternative"], result_df["CC"], color="#5680e9")
        ax.set_title("FTOPSIS Yakınlık Katsayıları (CC)")
        ax.set_ylabel("CC Değeri")
        fig.tight_layout()
        fig.savefig("ftopsis_grafik.png")

        # Son sonuçları kaydet
        self.last_result = {
            "fahp_ling": fahp_ling_experts,
            "ftopsis_ling": ftopsis_ling_experts,
            "weights": weights,
            "df": result_df
        }

        # İstatistik hesapla
        stats = statistical_summary(weights, result_df)
        
        # Uzman uyuşum analizi
        expert_agreement = expert_agreement_analysis(fahp_pfs_experts)
        
        # Radar chart
        radar_fig = create_radar_chart(self.criteria, weights)

        # Sonuç Penceresi
        win = tk.Toplevel(self.root)
        win.title("Sonuçlar")
        win.geometry("1000x850")
        win.configure(bg="#f2f2f2")

        tabs = ttk.Notebook(win)
        tabs.pack(expand=1, fill="both", padx=10, pady=10)

        # TAB 1 – FAHP Ağırlıkları
        tab1 = ttk.Frame(tabs)
        tabs.add(tab1, text="FAHP Ağırlıkları")

        tree = ttk.Treeview(tab1, columns=("Kriter", "Ağırlık"), show="headings", height=10)
        tree.heading("Kriter", text="Kriter")
        tree.heading("Ağırlık", text="Ağırlık")
        tree.column("Kriter", width=100)
        tree.column("Ağırlık", width=100)
        tree.pack(fill="both", expand=True, padx=10, pady=10)

        max_w = max(weights)
        pastel_green = "#d4f8d4"
        pastel_yellow = "#fff7c2"
        pastel_pink = "#ffd6e0"

        for i, w in enumerate(weights, 1):
            tag = f"w{i}"
            if w >= max_w * 0.85:
                color = pastel_green
            elif w >= max_w * 0.60:
                color = pastel_yellow
            else:
                color = pastel_pink
            tree.tag_configure(tag, background=color, foreground="black")
            tree.insert("", "end", values=(f"A{i}", f"{w:.4f}"), tags=(tag,))

        # TAB 2 – FTOPSIS Grafik
        tab2 = ttk.Frame(tabs)
        tabs.add(tab2, text="FTOPSIS Grafik")

        canvas = FigureCanvasTkAgg(fig, master=tab2)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

        # TAB 3 – Radar Chart
        tab3 = ttk.Frame(tabs)
        tabs.add(tab3, text="Radar Chart")

        canvas_radar = FigureCanvasTkAgg(radar_fig, master=tab3)
        canvas_radar.draw()
        canvas_radar.get_tk_widget().pack(fill="both", expand=True)

        # TAB 4 – İstatistik Özet
        tab4 = ttk.Frame(tabs)
        tabs.add(tab4, text="İstatistik Özet")

        txt_stats = tk.Text(tab4, font=("Segoe UI", 11), height=20)
        txt_stats.pack(fill="both", expand=True, padx=10, pady=10)
        
        txt_stats.insert("end", "=== İSTATİSTİKSEL ÖZET ===\n\n")
        for key, value in stats.items():
            txt_stats.insert("end", f"{key}: {value}\n")
        
        txt_stats.insert("end", "\n=== UZMAN UYUŞUM ANALİZİ ===\n\n")
        for key, value in expert_agreement.items():
            txt_stats.insert("end", f"{key}: {value}\n")
        
        txt_stats.configure(state="disabled")

        # TAB 5 – FTOPSIS Tablosu
        tab5 = ttk.Frame(tabs)
        tabs.add(tab5, text="FTOPSIS Tablosu")

        txt = tk.Text(tab5, font=("Segoe UI", 12))
        txt.pack(expand=True, fill="both")
        txt.insert("end", str(result_df))
        txt.configure(state="disabled")

        # PDF / Word butonları
        btn_frame = ttk.Frame(win)
        btn_frame.pack(pady=10)

        ttk.Button(
            btn_frame,
            text="PDF Raporu",
            command=lambda: self.ask_pdf_path(
                weights, result_df, fahp_ling_experts, pfs_fahp,
                ftopsis_ling_experts, pfs_top
            )
        ).grid(row=0, column=0, padx=10)

        ttk.Button(
            btn_frame,
            text="Word Raporu",
            command=lambda: self.ask_word_path(
                weights, result_df, fahp_ling_experts, pfs_fahp,
                ftopsis_ling_experts, pfs_top
            )
        ).grid(row=0, column=1, padx=10)

    # =================================================
    # 6. BÖLÜM: KAYDETME KONUMU & PDF/WORD RAPORLARI
    # =================================================
    def ask_pdf_path(self, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Dosyası", "*.pdf")],
            title="PDF Kaydetme Konumu Seç"
        )
        if file_path:
            self.create_pdf(file_path, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top)

    def ask_word_path(self, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Dosyası", "*.docx")],
            title="Word Kaydetme Konumu Seç"
        )
        if file_path:
            self.create_word(file_path, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top)

    def create_pdf(self, path, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top):
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Image
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors

        doc = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("<b>PFS–FAHP & PFS–FTOPSIS Raporu</b>", styles["Title"]))
        story.append(Spacer(1, 12))

        now = datetime.now().strftime("%d.%m.%Y %H:%M")
        story.append(Paragraph(f"Oluşturulma: {now}", styles["BodyText"]))
        story.append(Spacer(1, 12))

        story.append(Paragraph(
            "<b>Açıklama:</b> Güneş enerjisi yer seçimi için bulanık çok kriterli analiz.",
            styles["BodyText"]
        ))
        story.append(Spacer(1, 20))

        # FAHP – Uzman matrisleri
        story.append(Paragraph("<b>FAHP Uzman İkili Karşılaştırma Matrisleri</b>", styles["Heading2"]))
        for idx, exp in enumerate(fahp_ling_experts, 1):
            story.append(Paragraph(f"<b>Uzman {idx}</b>", styles["Heading3"]))
            data = [[""] + self.criteria]
            for i in range(5):
                row = [self.criteria[i]]
                for j in range(5):
                    row.append(exp[i][j])
                data.append(row)
            table = Table(data)
            table.setStyle([("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)])
            story.append(table)
            story.append(Spacer(1, 10))

        # FAHP – Aggregate PFS
        story.append(Paragraph("<b>FAHP Birleşik (Aggregate) PFS Matrisi</b>", styles["Heading2"]))
        data = [[""] + self.criteria]
        for i in range(5):
            row = [self.criteria[i]]
            for j in range(5):
                row.append(f"μ={pfs_fahp[i][j].mu:.3f}, ν={pfs_fahp[i][j].nu:.3f}")
            data.append(row)
        table = Table(data)
        story.append(table)
        story.append(Spacer(1, 15))

        # FAHP – Ağırlıklar
        story.append(Paragraph("<b>FAHP Kriter Ağırlıkları</b>", styles["Heading2"]))
        data = [["Kriter", "Ağırlık"]] + [[f"A{i+1}", f"{w:.4f}"] for i, w in enumerate(weights)]
        table = Table(data)
        table.setStyle([("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)])
        story.append(table)
        story.append(Spacer(1, 20))

        # FTOPSIS – Uzman matrisleri
        story.append(Paragraph("<b>FTOPSIS Uzman Değerlendirme Matrisleri</b>", styles["Heading2"]))
        for idx, exp in enumerate(ftopsis_ling_experts, 1):
            story.append(Paragraph(f"<b>Uzman {idx}</b>", styles["Heading3"]))
            data = [[""] + self.criteria]
            for i in range(5):
                row = [self.alternatives[i]]
                for j in range(5):
                    row.append(exp[i][j])
                data.append(row)
            table = Table(data)
            story.append(table)
            story.append(Spacer(1, 10))

        # FTOPSIS – Aggregate PFS
        story.append(Paragraph("<b>FTOPSIS Birleşik (Aggregate) PFS Matrisi</b>", styles["Heading2"]))
        data = [[""] + self.criteria]
        for i in range(5):
            row = [self.alternatives[i]]
            for j in range(5):
                row.append(f"μ={pfs_top[i][j].mu:.3f}, ν={pfs_top[i][j].nu:.3f}")
            data.append(row)
        table = Table(data)
        story.append(table)
        story.append(Spacer(1, 15))

        # FTOPSIS – Sonuç tablosu
        story.append(Paragraph("<b>FTOPSIS Sonuç Tablosu</b>", styles["Heading2"]))
        data2 = [df.columns.tolist()] + df.values.tolist()
        table2 = Table(data2)
        table2.setStyle([("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)])
        story.append(table2)
        story.append(Spacer(1, 15))

        # Grafik
        story.append(Paragraph("<b>FTOPSIS Grafiği</b>", styles["Heading2"]))
        story.append(Image("ftopsis_grafik.png", width=400, height=300))
        story.append(Spacer(1, 15))

        # Özet sonuç
        best_alt = df.iloc[0]["Alternative"]
        best_cc = df.iloc[0]["CC"]
        story.append(Paragraph(
            f"<b>En iyi alternatif:</b> {best_alt} (CC = {best_cc:.4f})",
            styles["BodyText"]
        ))

        doc.build(story)
        messagebox.showinfo("PDF", f"PDF raporu oluşturuldu:\n{path}")

    def create_word(self, path, weights, df, fahp_ling_experts, pfs_fahp, ftopsis_ling_experts, pfs_top):
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading("PFS–FAHP & PFS–FTOPSIS Raporu", 0)

        now = datetime.now().strftime("%d.%m.%Y %H:%M")
        doc.add_paragraph(f"Oluşturulma: {now}")

        doc.add_heading("Açıklama", level=1)
        doc.add_paragraph("Güneş enerjisi yer seçimi için bulanık çok kriterli karar verme analizi.")

        # FAHP – Uzman matrisleri
        doc.add_heading("FAHP Uzman İkili Karşılaştırma Matrisleri", level=1)
        for idx, exp in enumerate(fahp_ling_experts, 1):
            doc.add_heading(f"Uzman {idx}", level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            hdr[0].text = ""
            for i in range(5):
                hdr[i+1].text = self.criteria[i]
            for r in range(5):
                row_cells = table.add_row().cells
                row_cells[0].text = self.criteria[r]
                for c in range(5):
                    row_cells[c+1].text = exp[r][c]

        # FAHP – Aggregate PFS
        doc.add_heading("FAHP Birleşik PFS Matrisi", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = ""
        for i in range(5):
            hdr[i+1].text = self.criteria[i]
        for r in range(5):
            row_cells = table.add_row().cells
            row_cells[0].text = self.criteria[r]
            for c in range(5):
                row_cells[c+1].text = f"μ={pfs_fahp[r][c].mu:.3f}, ν={pfs_fahp[r][c].nu:.3f}"

        # FAHP – Ağırlıklar
        doc.add_heading("FAHP Ağırlıkları", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Kriter"
        hdr[1].text = "Ağırlık"
        for i, w in enumerate(weights):
            row = table.add_row().cells
            row[0].text = f"A{i+1}"
            row[1].text = f"{w:.4f}"

        # FTOPSIS – Uzman matrisleri
        doc.add_heading("FTOPSIS Uzman Değerlendirmeleri", level=1)
        for idx, exp in enumerate(ftopsis_ling_experts, 1):
            doc.add_heading(f"Uzman {idx}", level=2)
            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            hdr[0].text = ""
            for i in range(5):
                hdr[i+1].text = self.criteria[i]
            for r in range(5):
                row = table.add_row().cells
                row[0].text = self.alternatives[r]
                for c in range(5):
                    row[c+1].text = exp[r][c]

        # FTOPSIS – Aggregate PFS
        doc.add_heading("FTOPSIS Birleşik PFS Matrisi", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = ""
        for i in range(5):
            hdr[i+1].text = self.criteria[i]
        for r in range(5):
            row = table.add_row().cells
            row[0].text = self.alternatives[r]
            for c in range(5):
                row[c+1].text = f"μ={pfs_top[r][c].mu:.3f}, ν={pfs_top[r][c].nu:.3f}"

        # FTOPSIS – Sonuç Tablosu
        doc.add_heading("FTOPSIS Sonuç Tablosu", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = col
        for _, row_series in df.iterrows():
            row = table.add_row().cells
            for i, val in enumerate(row_series):
                row[i].text = str(val)

        # Grafik
        doc.add_heading("FTOPSIS Grafiği", level=1)
        doc.add_picture("ftopsis_grafik.png", width=Inches(5))

        # Özet Sonuç
        best_alt = df.iloc[0]["Alternative"]
        best_cc = df.iloc[0]["CC"]
        doc.add_heading("Özet Sonuç", level=1)
        doc.add_paragraph(f"En iyi alternatif: {best_alt} (CC = {best_cc:.4f})")

        doc.save(path)
        messagebox.showinfo("WORD", f"Word raporu oluşturuldu:\n{path}")


# ===========================================================
# =============== 7. BÖLÜM: PROGRAMI ÇALIŞTIR ===============
# ===========================================================

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
